"""DOCX parser using python-docx for text extraction."""

from pathlib import Path
from typing import Any

import structlog
from docx import Document

from src.ingestion_parsing.models.parsing_result import ParsingResult
from src.ingestion_parsing.parsers.base import BaseParser

logger = structlog.get_logger(__name__)


class DOCXParser(BaseParser):
    """DOCX document parser using python-docx library.
    
    Extracts text from:
    - Paragraphs
    - Tables
    - Headings
    - Lists
    
    Example:
        >>> parser = DOCXParser()
        >>> result = parser.parse("/path/to/document.docx")
        >>> print(result.text_content)
    """

    def parse(self, file_path: str) -> ParsingResult:
        """Parse DOCX document and extract all content and metadata.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ParsingResult with text, metadata, and structural information
            
        Raises:
            FileNotFoundError: If file does not exist
            ValueError: If DOCX parsing fails
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info("Starting DOCX parsing", file_path=file_path)
        
        try:
            # Extract text content
            text_content = self.extract_text(file_path)
            
            # Extract metadata
            metadata = self.extract_metadata(file_path)
            
            # Build parsing result
            result = ParsingResult(
                document_id=0,  # Will be set by calling service
                text_content=text_content,
                page_count=None,  # DOCX doesn't have fixed pages
                language=None,  # Could be detected with langdetect
                metadata=metadata,
                has_tables=metadata.get("has_tables", False),
                has_images=metadata.get("has_images", False),
                sections=metadata.get("sections", []),
            )
            
            logger.info(
                "DOCX parsing completed",
                file_path=file_path,
                text_length=len(text_content),
                has_tables=result.has_tables,
            )
            
            return result
            
        except Exception as e:
            logger.error(
                "DOCX parsing failed",
                file_path=file_path,
                error=str(e),
                error_type=type(e).__name__,
            )
            raise ValueError(f"Failed to parse DOCX: {e}") from e

    def extract_text(self, file_path: str) -> str:
        """Extract text content from DOCX.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            FileNotFoundError: If file does not exist
            ValueError: If text extraction fails
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Preserve heading structure
                    if paragraph.style.name.startswith('Heading'):
                        text_parts.append(f"\n# {paragraph.text}\n")
                    else:
                        text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            text_content = "\n\n".join(text_parts)
            
            logger.debug(
                "Text extraction completed",
                file_path=file_path,
                text_length=len(text_content),
            )
            
            return text_content
            
        except Exception as e:
            logger.error(
                "Text extraction failed",
                file_path=file_path,
                error=str(e),
            )
            raise ValueError(f"Failed to extract text: {e}") from e

    def extract_metadata(self, file_path: str) -> dict[str, Any]:
        """Extract metadata from DOCX.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary containing metadata (author, title, etc.)
            
        Raises:
            FileNotFoundError: If file does not exist
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract core properties
            core_props = doc.core_properties
            
            metadata: dict[str, Any] = {
                "author": core_props.author if core_props.author else None,
                "title": core_props.title if core_props.title else None,
                "subject": core_props.subject if core_props.subject else None,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "has_tables": len(doc.tables) > 0,
                "has_images": len(doc.inline_shapes) > 0,
                "sections": self._extract_sections(doc),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
            
            logger.debug("Metadata extraction completed", file_path=file_path, metadata=metadata)
            
            return metadata
            
        except Exception as e:
            logger.error("Metadata extraction failed", file_path=file_path, error=str(e))
            return {
                "has_tables": False,
                "has_images": False,
                "sections": [],
            }

    def supports_format(self, mime_type: str) -> bool:
        """Check if this parser supports DOCX format.
        
        Args:
            mime_type: MIME type to check
            
        Returns:
            True if MIME type is DOCX
        """
        return mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def _extract_sections(self, doc: Document) -> list[str]:
        """Extract section headings from document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of section heading texts
        """
        sections = []
        
        try:
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    sections.append(paragraph.text)
        except Exception as e:
            logger.warning("Section extraction failed", error=str(e))
        
        return sections
